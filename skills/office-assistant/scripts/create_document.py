#!/usr/bin/env python3
"""Office Assistant - Document Creator for Word, Excel, PPT."""

import argparse
import json
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import pandas as pd


def create_report(data, output):
    """Create a work report (工作汇报)."""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(12)
    
    # Title
    title = doc.add_heading(data.get('title', '工作汇报'), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"汇报人：{data.get('author', '')}    日期：{data.get('date', datetime.now().strftime('%Y年%m月%d日'))}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()
    
    # Overview section
    if 'overview' in data:
        doc.add_heading('一、工作概述', level=1)
        doc.add_paragraph(data['overview'])
    
    # Achievements section
    if 'achievements' in data:
        doc.add_heading('二、工作成果', level=1)
        for i, item in enumerate(data['achievements'], 1):
            p = doc.add_paragraph()
            run = p.add_run(f'{i}. ')
            run.bold = True
            p.add_run(item)
    
    # Data section
    if 'data' in data:
        doc.add_heading('三、数据指标', level=1)
        table = doc.add_table(rows=1, cols=len(data['data'].get('headers', [])))
        table.style = 'Light Grid Accent 1'
        
        # Headers
        for i, header in enumerate(data['data']['headers']):
            table.rows[0].cells[i].text = header
        
        # Data rows
        for row_data in data['data'].get('rows', []):
            row = table.add_row()
            for i, cell_data in enumerate(row_data):
                row.cells[i].text = str(cell_data)
    
    # Issues section
    if 'issues' in data:
        doc.add_heading('四、存在问题', level=1)
        for item in data['issues']:
            p = doc.add_paragraph(item, style='List Bullet')
    
    # Plans section
    if 'plans' in data:
        doc.add_heading('五、下一步计划', level=1)
        for i, item in enumerate(data['plans'], 1):
            p = doc.add_paragraph()
            run = p.add_run(f'{i}. ')
            run.bold = True
            p.add_run(item)
    
    # Signature
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f'汇报人：{data.get("author", "")}')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f'\n日期：{data.get("date", datetime.now().strftime("%Y年%m月%d日"))}')
    
    doc.save(output)
    print(f"✅ 工作汇报已保存: {output}")


def create_proposal(data, output):
    """Create a project proposal (项目方案)."""
    doc = Document()
    
    # Title
    title = doc.add_heading(data.get('title', '项目方案'), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Background
    if 'background' in data:
        doc.add_heading('一、项目背景', level=1)
        doc.add_paragraph(data['background'])
    
    # Objectives
    if 'objectives' in data:
        doc.add_heading('二、项目目标', level=1)
        for item in data['objectives']:
            doc.add_paragraph(item, style='List Bullet')
    
    # Scope
    if 'scope' in data:
        doc.add_heading('三、项目范围', level=1)
        doc.add_paragraph(data['scope'])
    
    # Timeline
    if 'timeline' in data:
        doc.add_heading('四、实施计划', level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        headers = ['阶段', '时间', '主要工作']
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for item in data['timeline']:
            row = table.add_row()
            row.cells[0].text = item.get('phase', '')
            row.cells[1].text = item.get('time', '')
            row.cells[2].text = item.get('work', '')
    
    # Budget
    if 'budget' in data:
        doc.add_heading('五、项目预算', level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        headers = ['项目', '金额（元）', '说明']
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for item in data['budget']:
            row = table.add_row()
            row.cells[0].text = item.get('item', '')
            row.cells[1].text = str(item.get('amount', ''))
            row.cells[2].text = item.get('note', '')
    
    # Risk
    if 'risks' in data:
        doc.add_heading('六、风险评估', level=1)
        for item in data['risks']:
            p = doc.add_paragraph()
            run = p.add_run(f"• {item.get('risk', '')}")
            run.bold = True
            p.add_run(f" - 应对措施：{item.get('mitigation', '')}")
    
    # Conclusion
    if 'conclusion' in data:
        doc.add_heading('七、结论', level=1)
        doc.add_paragraph(data['conclusion'])
    
    doc.save(output)
    print(f"✅ 项目方案已保存: {output}")


def create_summary(data, output):
    """Create a work summary (工作总结)."""
    doc = Document()
    
    # Title
    title = doc.add_heading(data.get('title', '工作总结'), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Period
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"总结期间：{data.get('period', '')}")
    
    doc.add_paragraph()
    
    # Main content sections
    sections = [
        ('overview', '一、总体概述'),
        ('achievements', '二、主要成绩'),
        ('highlights', '三、工作亮点'),
        ('problems', '四、存在不足'),
        ('lessons', '五、经验教训'),
        ('plans', '六、未来计划'),
    ]
    
    for key, title_text in sections:
        if key in data:
            doc.add_heading(title_text, level=1)
            content = data[key]
            if isinstance(content, list):
                for item in content:
                    doc.add_paragraph(item, style='List Bullet')
            else:
                doc.add_paragraph(content)
    
    doc.save(output)
    print(f"✅ 工作总结已保存: {output}")


def create_memo(data, output):
    """Create a memo (通知/公文)."""
    doc = Document()
    
    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(data.get('company', '公司名称'))
    run.bold = True
    run.font.size = Pt(16)
    
    # Document number
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"{data.get('doc_type', '通知')}〔{data.get('year', '2026')}〕{data.get('number', '001')}号")
    
    doc.add_paragraph()
    
    # Title
    title = doc.add_heading(data.get('title', ''), level=2)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Body
    doc.add_paragraph(data.get('content', ''))
    
    # Recipients
    if 'recipients' in data:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run('收：').bold = True
        p.add_run(', '.join(data['recipients']))
    
    # Sender
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(data.get('sender', ''))
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f"\n{data.get('date', datetime.now().strftime('%Y年%m月%d日'))}")
    
    doc.save(output)
    print(f"✅ 公文已保存: {output}")


def create_meeting_minutes(data, output):
    """Create meeting minutes (会议纪要)."""
    doc = Document()
    
    # Title
    title = doc.add_heading(data.get('title', '会议纪要'), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Meeting info
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.LEFT
    info.add_run(f"会议时间：{data.get('time', '')}\n")
    info.add_run(f"会议地点：{data.get('location', '')}\n")
    info.add_run(f"主 持 人：{data.get('chairman', '')}\n")
    info.add_run(f"参会人员：{', '.join(data.get('attendees', []))}")
    
    doc.add_paragraph()
    
    # Agenda
    if 'agenda' in data:
        doc.add_heading('一、会议议程', level=1)
        for i, item in enumerate(data['agenda'], 1):
            doc.add_paragraph(f"{i}. {item}")
    
    # Discussion
    if 'discussion' in data:
        doc.add_heading('二、讨论内容', level=1)
        for item in data['discussion']:
            p = doc.add_paragraph()
            run = p.add_run(f"• {item.get('topic', '')}")
            run.bold = True
            p.add_run(f"\n  {item.get('content', '')}")
    
    # Decisions
    if 'decisions' in data:
        doc.add_heading('三、会议决议', level=1)
        for i, item in enumerate(data['decisions'], 1):
            doc.add_paragraph(f"{i}. {item}")
    
    # Action items
    if 'actions' in data:
        doc.add_heading('四、待办事项', level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        headers = ['事项', '负责人', '截止时间', '状态']
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for item in data['actions']:
            row = table.add_row()
            row.cells[0].text = item.get('task', '')
            row.cells[1].text = item.get('owner', '')
            row.cells[2].text = item.get('deadline', '')
            row.cells[3].text = item.get('status', '待完成')
    
    doc.save(output)
    print(f"✅ 会议纪要已保存: {output}")


def main():
    parser = argparse.ArgumentParser(description="Office Document Creator")
    parser.add_argument("--type", "-t", required=True,
                       choices=['report', 'proposal', 'summary', 'memo', 'minutes'],
                       help="Document type")
    parser.add_argument("--data", "-d", required=True, help="JSON data for document")
    parser.add_argument("--output", "-o", required=True, help="Output file path")
    
    args = parser.parse_args()
    
    data = json.loads(args.data) if isinstance(args.data, str) else args.data
    
    creators = {
        'report': create_report,
        'proposal': create_proposal,
        'summary': create_summary,
        'memo': create_memo,
        'minutes': create_meeting_minutes,
    }
    
    creator = creators.get(args.type)
    if creator:
        creator(data, args.output)
    else:
        print(f"❌ Unknown document type: {args.type}")


if __name__ == "__main__":
    main()
